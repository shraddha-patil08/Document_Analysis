import PyPDF2
import streamlit as st
import PyPDF2
import docx
import os
from transformers import pipeline
from groq import Groq

api_key = os.getenv("GROQ_API_KEY")

# API Configuration (Replace with your actual API key)
client = Groq(api_key=api_key)


def extract_title_and_text(uploaded_file, file_type):
    title = "Untitled Document"
    text = ""

    if file_type == "pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        metadata = reader.metadata

        if metadata and metadata.title:
            title = metadata.title.strip()
        else:
            first_page_text = reader.pages[0].extract_text() or ""
            lines = first_page_text.split("\n")
            for line in lines:
                if line.strip():
                    title = line.strip()
                    break

        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"

        if not text.strip():
            text = "No text found in PDF."

    elif file_type == "docx":
        doc = docx.Document(uploaded_file)
        if doc.paragraphs:
            title = doc.paragraphs[0].text.strip() or "Untitled Document"
        text = "\n".join([para.text for para in doc.paragraphs]) or "No text found in DOCX."

    elif file_type == "txt":
        content = uploaded_file.read().decode("utf-8", errors="ignore").strip()
        title = content.split("\n")[0] if content else "Untitled Document"
        text = content or "No text found in TXT."

    return title, text


# feature extracrion
def extract_features(uploaded_file, file_type):
    metadata = {
        "Title": "Untitled Document",
        "Author": "Unknown",
        "Creation Date": "N/A",
        "Headings": [],
        "Paragraphs": [],
        "Tables": 0,
        "Images": 0
    }
    text = ""

    if file_type == "pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        pdf_metadata = reader.metadata

        if pdf_metadata:
            metadata["Title"] = pdf_metadata.title or "Untitled Document"
            metadata["Author"] = pdf_metadata.author or "Unknown"
            metadata["Creation Date"] = pdf_metadata.get("/CreationDate", "N/A")

        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
            lines = page_text.split("\n")

            for line in lines:
                if len(line) < 50 and line.isupper():
                    metadata["Headings"].append(line)

    elif file_type == "docx":
        doc = docx.Document(uploaded_file)
        metadata["Title"] = doc.paragraphs[0].text.strip() if doc.paragraphs else "Untitled Document"

        for para in doc.paragraphs:
            text += para.text + "\n"
            if para.style.name.startswith("Heading"):
                metadata["Headings"].append(para.text)
            else:
                metadata["Paragraphs"].append(para.text)

        metadata["Tables"] = len(doc.tables)
        metadata["Images"] = sum(1 for rel in doc.part.rels if "image" in doc.part.rels[rel].target_ref)

    elif file_type == "txt":
        content = uploaded_file.read().decode("utf-8", errors="ignore").strip()
        text = content
        metadata["Title"] = content.split("\n")[0] if content else "Untitled Document"
        metadata["Paragraphs"] = content.split("\n\n")

    return metadata, text.strip()



# key element extraction
def extract_key_elements(text):
    if not text.strip():
        return "No text available for analysis."

    prompt = f"""
    Extract the key elements from the following document in pointwise manner:
    
    - **Key Findings**
    - **Conclusions**
    - **Recommendations**
    - **Important Terms & Sentences**
    
    Text:
    {text}
    """

    chat_completion = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )

    return chat_completion.choices[0].message.content.strip()